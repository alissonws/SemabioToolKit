#!/usr/bin/env python
# -*- coding: utf-8 -*-

#from gimpfu import *
import os, subprocess, platform, sys, ctypes, logging

from openpyxl import load_workbook
from docx import Document

if platform.system() == "Windows":
    import comtypes.client



class PyCertsSB:
    def __init__(self):
        pass

    def check_file(self, dir,what,headers=0,col_args=[]):
        """
        Método para checar a viabilidade da planilha e/ou contar o número de certificados.
        Sem os dois últimos parâmetros, faz o primeiro, com todos os parâmetros preenchidos, faz o segundo.
        :param dir: diretório da planilha;
        :param col_args: tuple com colunas de dados referente aos certificados;
        :param headers: considerar a primeira linha como cabeçalho da planilha;
        :return:
        """
        if what == ".xlsx":
            try:
                wb = load_workbook(dir)
                ws = wb.active

                if col_args:
                    mount = len(ws[col_args[0][0]]) - headers

                    return mount
            except Exception as e:
                print(e)
                raise ImportError("Planilha invállida")
        elif what == ".docx":
            try:
                preview = Document(dir)
            except:
                raise ImportError("Arquivo invállido")



    def generateDocx(self,template_plan, plan_path,save_plan, headers,col_args,write):
        wb = load_workbook(plan_path)
        ws = wb.active


        lim = len(ws[col_args[0][0]]) - headers
        logging.debug("There were found {} lines in this sheet".format(lim))

        start = 1 + headers
        count = 0


        docx_list = []


        while count < lim:
            skip = 0

            template = open(template_plan, "rb")

            template_file = Document(template)

            for tuple in col_args:
                cell = tuple[0]
                tag = tuple[1]


                if ws[cell+str(start+count)].value == None:
                    skip = 1
                    logging.info("The cell is empty")
                    break

                for paragraph in template_file.paragraphs:
                    for run in paragraph.runs:
                        if tag in run.text:
                            run.text = run.text.replace(tag,str(ws[cell+str(start+count)].value))

            if skip == 1:
                break
            file = (save_plan+"/"+str(start+count)+".docx")

            if platform.system() == "Windows":
                file = file.replace("/","\\")

            template_file.save(file)
            template.close()

            docx_list.append(file) #Por algum motivo eu to salvando o nome dos arquivos mas n sei pq eu fiz isso então só vou deixar aqui pq não machuca ninguém

            subprocess.run('"C:\Program Files\LibreOffice\program\soffice" --convert-to pdf --outdir "{}" "{}"'.format(save_plan, file))

            os.remove(file)

            if write != 0:
                ws[write+str(start+count)].value = str(start+count)+".pdf"

            logging.debug("File named {} was generated successfully at {}".format((str(start+count)+".pdf"),save_plan))
            count += 1

        if write != 0:
            wb.save(plan_path)

        logging.debug("Generation completed")